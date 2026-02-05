# Note: Install the python-docx library before running this script
# Run in terminal: pip install python-docx
# Also install matplotlib: pip install matplotlib

import os
from pathlib import Path
import docx
import matplotlib.pyplot as plt

def count_words_in_docx(file_path):
    try:
        doc = docx.Document(file_path)
        word_count = 0

        # Count words in paragraphs
        for para in doc.paragraphs:
            words = para.text.split()
            word_count += len(words)

        # Count words in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Avoid counting the same cell multiple times (in case of merged cells)
                    if cell.text.strip():  # Only process non-empty cells
                        words = cell.text.split()
                        word_count += len(words)

        return word_count
    except Exception as e:
        print(f"Error processing {file_path}: {str(e)}")
        return 0

def count_words_in_directory(directory_path):
    directory = Path(directory_path)
    if not directory.is_dir():
        print(f"Error: {directory_path} is not a valid directory")
        return []
    
    docx_files = list(directory.glob("*.docx"))
    
    if not docx_files:
        print("No .docx files found in the directory")
        return []
    
    word_counts = []
    for file_path in docx_files:
        words = count_words_in_docx(file_path)
        print(f"{file_path.name}: {words} words")
        word_counts.append((file_path.name, words))
    
    return word_counts

def plot_word_distribution(categorized_counts, total_words):
    all_items = []
    for category, counts in categorized_counts:
        for name, count in counts:
            all_items.append((name, count, category))
    
    # Sort by word count descending
    all_items.sort(key=lambda x: x[1], reverse=True)
    
    # Prepare data for plotting
    filenames = [item[0].replace('.docx', '').replace(' Notes', '') for item in all_items]
    counts = [item[1] for item in all_items]
    categories = [item[2] for item in all_items]
    
    # Assign colors based on category
    color_map = {
        'Books/Podcasts': 'skyblue',
        'Learning': 'lightgreen',
        'Projects': 'salmon',
        'Classes': 'gold'
    }
    colors = [color_map[cat] for cat in categories]
    
    # Create bar chart
    plt.figure(figsize=(12, 8))
    plt.bar(filenames, counts, color=colors)
    plt.xlabel('Document')
    plt.ylabel('Word Count')
    plt.title('Word Count Distribution by Document')
    plt.xticks(rotation=45, ha='right', fontsize=9)
    
    # Add legend
    handles = [plt.Rectangle((0,0),1,1, color=color) for color in color_map.values()]
    plt.legend(handles, color_map.keys())
    
    # Add total word count text
    plt.text(0.5, 0.94, f'Total Words: {total_words}', ha='center', va='bottom', transform=plt.gca().transAxes, fontsize=10)
    
    plt.tight_layout()
    plt.savefig('word_count_distribution.png', dpi=300, bbox_inches='tight')
    plt.show()

if __name__ == "__main__":
    counts_books_podcasts = count_words_in_directory("Notes Books - Podcasts/")
    counts_learning = count_words_in_directory("Notes Learning/")
    counts_projects = count_words_in_directory("Notes Projects/")
    counts_classes = count_words_in_directory("Notes Classes/")
    
    categorized_counts = [
        ('Books/Podcasts', counts_books_podcasts),
        ('Learning', counts_learning),
        ('Projects', counts_projects),
        ('Classes', counts_classes)
    ]
    
    # Calculate total words
    total_words = sum(count for _, sublist in categorized_counts for _, count in sublist)
    print(f"\nTotal words in all .docx files: {total_words}")
    
    # Plot the distribution
    plot_word_distribution(categorized_counts, total_words)